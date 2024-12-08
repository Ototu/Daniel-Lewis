#AOA Project
#Date November 18,2024
#: Khadejah Benjamin-2208656 
# Ramon Johnston -2008317
#Daniel Lewis- 2202361 
# Rushane Green – 2006930 
# Chamarie Taylor – 2100037

import shutil
import tkinter as tk
from tkinterdnd2 import TkinterDnD, DND_FILES
import os
from tkinter import filedialog, messagebox, simpledialog, ttk
from datetime import datetime
from tkinter import font
import threading
import pytesseract
import mimetypes
from PIL import Image as PILImage, ImageTk
from PyPDF2 import PdfReader
import PyPDF2
from docx import Document
import openai

openai.api_key = " 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"  # Replace with your actual API key
pytesseract.pytesseract.tesseract_cmd = r'C:\Users\lewis\AppData\Local\Programs\Tesseract-OCR\tesseract.exe' 

class FileBrowserApp:

    def __init__(self, root):
        self.root = root
        self.root.title("File Browser")
        self.root.geometry("800x600")
        self.root.configure(bg="#f0f0f0")
        self.root.iconbitmap(self.get_icon_path('aoa_group_27V_icon.ico'))  # Use dynamic icon path

        # Load icons/images with dynamic paths
        self.new_folder_icon = self.load_icon('folder.png')
        self.new_file_icon = self.load_icon('fil.png')
        self.new_search_icon = self.load_icon('search.png')
        self.new_upload_icon = self.load_icon('upload.png')
        self.sort_icon = self.load_icon('list.png')

        # Side panel frame
        self.sidebar = tk.Frame(self.root, width=150, bg="#001540")
        self.sidebar.pack(side=tk.LEFT, fill=tk.Y)
        
        # New Folder Button
        new_folder_btn = tk.Button(self.sidebar, image=self.new_folder_icon, compound=tk.LEFT, text='New Folder', fg='white', borderwidth=0, highlightthickness=0, bg="#001540", activebackground="#001540", command=self.create_new_folder)
        new_folder_btn.pack(pady=10, padx=10, fill=tk.X)

        # New File Button
        new_file_btn = tk.Button(self.sidebar, image=self.new_file_icon, compound=tk.LEFT, text='New File', fg='white', borderwidth=0, highlightthickness=0, bg="#001540", activebackground="#001540", command=self.create_new_file)
        new_file_btn.pack(pady=10, padx=8, fill=tk.X)

        # Top bar for search bar and upload
        self.top_bar = tk.Frame(self.root, bd=1, relief=tk.RAISED)
        self.top_bar.pack(side=tk.TOP, fill=tk.X)

        # Search bar
        self.search_entry = tk.Entry(self.top_bar, width=20)
        self.search_entry.pack(side=tk.LEFT, padx=5, pady=5)

        # Search button/icon
        search_button = tk.Button(self.top_bar, image=self.new_search_icon, compound=tk.LEFT, text="Search", command=self.search_file, borderwidth=0, highlightthickness=0) 
        search_button.pack(side=tk.LEFT, padx=5, pady=5)

        # Upload button/icon
        upload_button = tk.Button(self.top_bar, image=self.new_upload_icon, compound=tk.LEFT, text="Upload", command=self.open_folder, borderwidth=0, highlightthickness=0)  
        upload_button.pack(side=tk.RIGHT, padx=5, pady=5)

        # Sort button/icon (icon for dropdown)
        sort_button = tk.Button(self.top_bar, image=self.sort_icon, compound=tk.LEFT, command=self.show_sort_menu, borderwidth=0, highlightthickness=0, bg="#f0f0f0")  
        sort_button.pack(side=tk.RIGHT, padx=10, pady=5)

        # Dynamic path tracking and history for back navigation
        self.current_path = ""
        self.path_history = []

        # Breadcrumb Navigation
        self.breadcrumb_frame = tk.Label(self.root, text="Home", bg="#f0f0f0", anchor="w", padx=10)
        self.breadcrumb_frame.pack(fill=tk.X)

        # Treeview to display files
        self.treeview_frame = tk.Frame(self.root)
        self.treeview_frame.pack(fill=tk.BOTH, expand=1)

        self.file_tree = ttk.Treeview(self.treeview_frame, columns=("Name", "Extension", "Size", "Modified", "Content Summary"), show="headings")
        self.file_tree.heading("Name", text="File Name")
        self.file_tree.heading("Extension", text="Extension")
        self.file_tree.heading("Size", text="Size")
        self.file_tree.heading("Modified", text="Date Modified")
        self.file_tree.heading("Content Summary", text="Content Summary")

        self.file_tree.column("Name", width=300)
        self.file_tree.column("Extension", width=80, anchor="center")
        self.file_tree.column("Size", width=100, anchor="e")
        self.file_tree.column("Modified", width=150, anchor="center")
        self.file_tree.column("Content Summary", width=250)

        self.file_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=1)

        self.scrollbar = tk.Scrollbar(self.treeview_frame, orient=tk.VERTICAL, command=self.file_tree.yview)
        self.scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.file_tree.configure(yscrollcommand=self.scrollbar.set)


        # File info display
        self.info_label = tk.Label(self.root, text="", wraplength=600, justify=tk.LEFT, bg="#f0f0f0")
        self.info_label.pack(pady=10)

        # Preview area for file content
        self.preview_area = tk.Label(self.root, text="File Preview will appear here", bg="#f0f0f0", justify=tk.LEFT, anchor="w")
        self.preview_area.pack(pady=10, padx=10)

        # Bind events for double-click and selection
        self.file_tree.bind("<Double-Button-1>", self.open_file)
        self.file_tree.bind("<<TreeviewSelect>>", self.display_file_info)

        # Add a context menu for file actions (rename, delete)
        self.context_menu = tk.Menu(self.root, tearoff=0)
        self.context_menu.add_command(label="Rename", command=self.rename_file)
        self.context_menu.add_command(label="Delete", command=self.delete_file)
        self.file_tree.bind("<Button-3>", self.show_context_menu)
    

    def load_icon(self, icon_name):
        try:
            return tk.PhotoImage(file=self.get_icon_path(icon_name))
        except Exception as e:
            messagebox.showwarning("Warning", f"Icon {icon_name} not found. Using default icon.")
            return tk.PhotoImage()  # Return a blank icon if the specified one is missing
        
    def show_context_menu(self, event):
        item = self.file_tree.identify('item', event.x, event.y)
        if item:
            self.file_tree.selection_set(item)
            self.context_menu.post(event.x_root, event.y_root)

    def rename_file(self):
        selected_item = self.file_tree.selection()[0]
        current_name = self.file_tree.item(selected_item, "values")[0]
        file_path = os.path.join(self.current_path, current_name)

        
    
    # Get the file extension
        file_name, file_extension = os.path.splitext(current_name)
    
    # Ask user for the new name (without extension)
        new_name_without_ext = simpledialog.askstring("Rename File", f"Enter a new name for {file_name}:")
    
        if new_name_without_ext:
        # Ensure the new file name includes the old extension
         new_name = new_name_without_ext + file_extension
        new_file_path = os.path.join(self.current_path, new_name)
        
        try:
            os.rename(file_path, new_file_path)
            self.list_directory(self.current_path)
        except Exception as e:
            messagebox.showerror("Error", f"Failed to rename file: {e}")

    def update_summary(self, file_path, tree_item):
        try:
            #Generate Summary for the renamed file
            summary = self.generate_content_summary(file_path)

            #update the treeview item with the new summary
            self.file_tree.item(tree_item, values=(os.path.basename(file_path),self.file_tree.item(tree_item, "values")[1],#extension
                                                   self.file_tree.item(tree_item, "values")[2], # size
                                                    self.file_tree.item(tree_item, "values")[3], # date
                                                     summary )) # New Summary
        except Exception as e:
            messagebox.showerror("Error", f"Failed to update file summary: {e}")
         
    def delete_file(self):
        selected_item = self.file_tree.selection()[0]
        file_name = self.file_tree.item(selected_item, "values")[0]
        file_path = os.path.join(self.current_path, file_name)
        
        confirm = messagebox.askyesno("Delete File", f"Are you sure you want to delete {file_name}?")
        if confirm:
            try:
                if os.path.isdir(file_path):
                    shutil.rmtree(file_path)
                else:
                    os.remove(file_path)
                self.list_directory(self.current_path)
            except Exception as e:
                messagebox.showerror("Error", f"Failed to delete file: {e}")

    def update_breadcrumb(self):
        for widget in self.breadcrumb_frame.winfo_children():
            widget.destroy()

        path_parts = self.current_path.split(os.sep)
        path_accumulated = ""

        for idx, part in enumerate(path_parts):
            if part:
                path_accumulated = os.path.join(path_accumulated, part)
                part_label = tk.Label(
                    self.breadcrumb_frame,
                    text=part,
                    fg="blue",
                    bg="#f0f0f0",
                    cursor="hand2"
                )
                part_label.bind("<Button-1>", lambda e, p=path_accumulated: self.navigate_to_path(p))
                part_label.pack(side=tk.LEFT)

                if idx < len(path_parts) - 1:
                    separator = tk.Label(self.breadcrumb_frame, text=">", bg="#f0f0f0")
                    separator.pack(side=tk.LEFT)



    def navigate_to_path(self, path):
        if os.path.isdir(path):
            self.current_path = path
            self.list_directory(path)

    def format_file_size(self, size_bytes):
        if size_bytes < 1024:
            return f"{size_bytes} B"
        elif size_bytes < 1024 ** 2:
            return f"{size_bytes / 1024:.1f} KB"
        elif size_bytes < 1024 ** 3:
            return f"{size_bytes / (1024 ** 2):.1f} MB"
        else:
            return f"{size_bytes / (1024 ** 3):.1f} GB"

    def get_icon_path(self, icon_name):
        script_dir = os.path.dirname(os.path.realpath(__file__))
        return os.path.join(script_dir, icon_name)

    def open_folder(self):
        folder_path = filedialog.askdirectory()
        if folder_path:
            self.path_history.append(self.current_path)
            self.current_path = folder_path
            self.list_directory(folder_path)
            self.process_files_for_summary(folder_path)
            

    def process_files_for_summary(self, folder_path):
        try:
            files_info = self.get_files_info(folder_path)

            for file_info in files_info:
                file_name = file_info[0]
                file_path = os.path.join(folder_path, file_name)

                if file_name.lower().endswith(".txt"):
                  with open(file_path, 'r') as file:
                    content = file.read()
                    summary = self.summarize_text_with_openai(content)
                    print(f"Summary for {file_name}: {summary}")

                elif file_name.lower().endswith(".pdf"):
                   summary = self.summarize_pdf(file_path)
                   print(f"Summary for {file_name}: {summary}")
  
                elif file_name.lower().endswith(".docx"):
                   summary = self.summarize_word(file_path)
                   print(f"Summary for {file_name}: {summary}")

                elif file_name.lower().endswith((".png", ".jpg", ".jpeg", ".gif")):
                  summary = self.summarize_image(file_path)
                  print(f"Summary for {file_name}: {summary}")
                  
                  for item in self.file_tree.get_children():
                   if self.file_tree.item(item, "values")[0] == file_name:
                    self.file_tree.item(item, values=(file_name, file_info[1], file_info[2], file_info[3], summary))
                    break
             
        except Exception as e:
          messagebox.showerror("Error", f"Failed to process folder for summarization: {e}")

    def list_directory(self, path):
        try:
            self.file_tree.delete(*self.file_tree.get_children()) # Clear Current Treeview
            self.current_path = path
            files_info = self.get_files_info(path)

            for file_info in files_info:
                file_name= file_info[0]
                file_path=os.path.join(path, file_name)

                summary= self.generate_content_summary(file_path)
                #regenerate and include the summary for the file
                self.file_tree.insert("", tk.END, values=(file_name, file_info[1], file_info[2], file_info[3], summary))

            self.update_breadcrumb()

        except Exception as e:
            messagebox.showerror("Error", f"Failed to list directory: {e}")

    def get_files_info(self, path):
        files_info = []
        for file_name in os.listdir(path):
            file_path = os.path.join(path, file_name)
            if os.path.isdir(file_path):
                file_size = "-"
                modified_time = datetime.fromtimestamp(os.path.getmtime(file_path)).strftime("%Y-%m-%d %H:%M:%S")
                extension = "-"
                content_summary="-"
            else:
                raw_file_size = os.path.getsize(file_path)
                file_size = self.format_file_size(raw_file_size)
                modified_time = datetime.fromtimestamp(os.path.getmtime(file_path)).strftime("%Y-%m-%d %H:%M:%S")
                extension = file_name.split('.')[-1] if '.' in file_name else "-"
                content_summary="Summary not Available"

            files_info.append((file_name, extension, file_size, modified_time, content_summary))

        return files_info

    def open_file(self):
        selected_item = self.file_tree.focus()
        file_name = self.file_tree.item(selected_item, "values")[0]
        file_path = os.path.join(self.current_path, file_name)

        try:
            if file_name.endswith(".txt"):
                with open(file_path, 'r') as file:
                    content = file.read()
                    self.preview_area.config(text=content[:500])  # Preview the first 500 characters
            elif file_name.endswith(".pdf"):
                with open(file_path, "rb") as file:
                    reader = PdfReader(file)
                    content = ""
                    for page in reader.pages[:2]:  # Limit to first 2 pages
                        content += page.extract_text()
                    self.preview_area.config(text=content[:500])
            elif file_name.endswith(".docx"):
                doc = Document(file_path)
                content = "\n".join([p.text for p in doc.paragraphs[:2]])  # Preview first 2 paragraphs
                self.preview_area.config(text=content[:500])
            elif file_name.lower().endswith((".png", ".jpg", ".jpeg", ".gif")):
                image = PILImage.open(file_path)
                image.thumbnail((300, 300))
                photo = ImageTk.PhotoImage(image)
                self.preview_area.config(image=photo)
                self.preview_area.image = photo  # Keep reference to image object
            else:
                self.preview_area.config(text="File format not supported for preview.")
        except Exception as e:
            self.preview_area.config(text=f"Error opening file: {e}")

    def display_file_info(self, event):
        # Get the selected item in the Treeview
        selected_item = self.file_tree.selection()[0]
        file_name = self.file_tree.item(selected_item, "values")[0]
        file_path = os.path.join(self.current_path, file_name)
        
        try:
            # Get file stats
            file_info = os.stat(file_path)
            summary = self.generate_content_summary(file_path)
            file_size = self.format_file_size(file_info.st_size)
            modified_time = datetime.fromtimestamp(file_info.st_mtime).strftime("%Y-%m-%d %H:%M:%S")
            self.info_label.config(text=f"File: {file_name}\nSize: {file_size}\nLast Modified: {modified_time}\nSummary: {summary}")
            self.preview_file(file_path)

        except Exception as e:
            messagebox.showerror("Error", f"Failed to retrieve file info: {e}")

    def preview_file(self, file_path):
      file_type, _ = mimetypes.guess_type(file_path)
      if file_type:
        if file_type.startswith('text'):
            # Preview text files by reading the first 200 characters
            with open(file_path, 'r') as file:
                content = file.read(200)  # First 200 characters of text file
                self.preview_area.config(text=f"Preview:\n{content}")
        
        elif file_type.startswith('image'):
            # Preview image files by showing a thumbnail
            img = PILImage.open(file_path)
            img.thumbnail((100, 100))
            img = ImageTk.PhotoImage(img)
            self.preview_area.config(image=img, text="")  # Show image
            self.preview_area.image = img
        
        elif file_type == 'application/pdf':
            # Preview for PDF (placeholder for first page)
            self.preview_area.config(text="PDF Preview (First page will appear here)")  # Placeholder
        
        elif file_type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
            # Preview DOCX files by reading the first few paragraphs
            try:
                doc = Document(file_path)
                preview_text = "\n".join([p.text for p in doc.paragraphs[:3]])  # First 3 paragraphs
                self.preview_area.config(text=f"DOCX Preview:\n{preview_text}")
            except Exception as e:
                self.preview_area.config(text=f"Error reading DOCX: {str(e)}")
        else:
            self.preview_area.config(text="No preview available for this file type.")

    
    def process_files_for_summary(self, folder_path):
        try:
            files_info = self.get_files_info(folder_path)
            for file_info in files_info:
                file_name = file_info[0]
                file_path = os.path.join(folder_path, file_name)

                if not os.path.isdir(file_path):
                    summary = self.generate_content_summary(file_path)
                    for item in self.file_tree.get_children():
                        if self.file_tree.item(item, "values")[0] == file_name:
                            self.file_tree.item(item, values=(file_name, file_info[1], file_info[2], file_info[3], summary))
                            break
        except Exception as e:
            messagebox.showerror("Error", f"Failed to process folder for summarization: {e}")

    def generate_content_summary(self, file_path):
        try:
            if file_path.endswith(".txt"):
                with open(file_path, 'r') as file:
                    text = file.read()
            elif file_path.endswith(".pdf"):
                text = self.extract_text_from_pdf(file_path)
            elif file_path.endswith(".docx"):
                text = self.extract_text_from_word(file_path)
            elif file_path.lower().endswith((".jpg", ".jpeg", ".png")):
                text = self.extract_text_from_image(file_path)
            else:
                return "Unsupported file format for summary"

            # Summarize text using AI
            summary = self.summarize_text( [text])
            return summary
        except Exception as e:
            return f"Error summarizing content: {e}"
    

    def extract_text_from_pdf(self, file_path):
        reader = PyPDF2.PdfReader(file_path)
        return " ".join([page.extract_text() for page in reader.pages])

    def extract_text_from_word(self, file_path):
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])

    def extract_text_from_image(self, file_path):
        return pytesseract.image_to_string(PILImage.open(file_path))
    
    def summarize_text(self,text):
        try:
            # Split text into smaller chunks to avoid token limit issues
            text_chunks = [text[i:i+1000] for i in range(0, len(text), 1000)]
            summaries = []

            for chunk in text_chunks:
                # Make a request to OpenAI's API for summarization
                response = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",
                    messages=[ 
                        {"role": "system", "content": "You are a helpful research assistant."},
                        {"role": "user", "content": f"Summarize the following text: {chunk}"}
                    ]
                )

                page_summary = response['choices'][0]['message']['content']
                summaries.append(page_summary)

            return "\n".join(summaries)
        except Exception as e:
            return f"Error summarizing text: {e}"
        
    def create_new_folder(self):
        folder_name = simpledialog.askstring("Folder Name", "Enter new folder name:")
        if folder_name:
            new_folder_path = os.path.join(self.current_path, folder_name)
            try:
                os.mkdir(new_folder_path)
                self.list_directory(self.current_path)
            except Exception as e:
                messagebox.showerror("Error", f"Failed to create folder: {e}")

    def create_new_file(self):
        file_name = simpledialog.askstring("File Name", "Enter new file name:")
        if file_name:
            new_file_path = os.path.join(self.current_path, file_name)
            try:
                with open(new_file_path, 'w') as file:
                    file.write("")  # Empty content initially
                self.list_directory(self.current_path)
            except Exception as e:
                messagebox.showerror("Error", f"Failed to create file: {e}")


    def search_file(self):
        search_query = self.search_entry.get()
        if search_query:
            self.list_directory(self.current_path)
            self.search_in_files(search_query)
            
    def search_in_files(self, query):
        for item in self.file_tree.get_children():
            file_name = self.file_tree.item(item)["values"][0]
            if query.lower() in file_name.lower():
                self.file_tree.selection_set(item)

    def show_sort_menu(self):
        sort_menu = tk.Menu(self.root, tearoff=0)
        sort_menu.add_command(label="Sort by Name", command=lambda: self.sort_files("Name"))
        sort_menu.add_command(label="Sort by Size", command=lambda: self.sort_files("Size"))
        sort_menu.add_command(label="Sort by Modified Date", command=lambda: self.sort_files("Modified"))
        sort_menu.post(self.root.winfo_pointerx(), self.root.winfo_pointery())

    def sort_files(self, criteria):
        try:
            self.file_tree.delete(*self.file_tree.get_children())
            files_info = self.get_files_info(self.current_path)

            if criteria == "Name":
                files_info.sort(key=lambda x: x[0].lower())
            elif criteria == "Size":
                files_info.sort(key=lambda x: x[2])  # Sort by raw file size
            elif criteria == "Modified":
                files_info.sort(key=lambda x: x[3])  
                
            for item in files_info:
                file_name=item[0]
                file_path=os.path.join(self.current_path, file_name)

                summary= self.generate_content_summary(file_path)
                   

          
                self.file_tree.insert("", tk.END, values=(item[0],item[1],item[2],item[3], summary))
        except Exception as e:
            messagebox.showerror("Error", f"Failed to sort files: {e}")
    

# Start the Tkinter application
if __name__ == "__main__":
    root = tk.Tk()
    app = FileBrowserApp(root)
    root.mainloop()

